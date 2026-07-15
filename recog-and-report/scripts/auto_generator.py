import os
import re
import json
import argparse
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_run_font(run, font_name, font_size_pt, bold=False, italic=False):
    """Sets the East Asian and ASCII fonts for a run to support mixed layouts."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_cell_text(cell, text, font_name='宋体', font_size_pt=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Overwrites cell text while keeping alignment, line spacing, and fonts."""
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(text)
    set_run_font(run, font_name, font_size_pt, bold=bold)
    
    for extra_p in list(cell.paragraphs[1:]):
        extra_p._element.getparent().remove(extra_p._element)

def add_toc_entry(p_ref, text, page_num, indent_level=0):
    """Inserts a table of contents entry before the referenced paragraph."""
    p = p_ref.insert_paragraph_before()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if indent_level > 0:
        p.paragraph_format.left_indent = Pt(indent_level * 16)
        
    char_len = 0
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            char_len += 2
        else:
            char_len += 1
            
    dots_needed = 72 - char_len
    if dots_needed < 5:
        dots_needed = 5
    dots = "…" * (dots_needed // 2)
    
    is_bold = (indent_level == 0)
    
    run_text = p.add_run(text)
    set_run_font(run_text, '宋体', 12, bold=is_bold)
    
    run_dots = p.add_run(dots)
    set_run_font(run_dots, '宋体', 12, bold=False)
    
    run_num = p.add_run(str(page_num))
    set_run_font(run_num, '宋体', 12, bold=is_bold)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '黑体', 14, bold=True)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '黑体', 12, bold=True)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '楷体', 12, bold=False)
    return p

def add_body_paragraph(doc, text, indent_pt=24):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.first_line_indent = Pt(indent_pt)
    
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, bold=False)
    return p

def add_reference_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    
    run = p.add_run(text)
    set_run_font(run, '宋体', 10.5, bold=False)
    return p

def find_template_landmarks(doc):
    """
    Dynamically identifies:
    - cover_end_idx: Last cover page paragraph (e.g. date paragraph)
    - section_break_idx: First paragraph holding a section break (w:sectPr)
    """
    cover_end_idx = None
    section_break_idx = None
    
    date_pattern = re.compile(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日')
    
    for idx, p in enumerate(doc.paragraphs):
        # 1. Match date paragraph
        if date_pattern.search(p.text):
            cover_end_idx = idx
            
        # 2. Match section break paragraph
        pPr = p._p.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None and section_break_idx is None:
            section_break_idx = idx
            
    # Fallback to sensible defaults if landmarks are not found
    if cover_end_idx is None:
        cover_end_idx = 12
    if section_break_idx is None:
        # If no section break paragraph found, find the last paragraph index before index 65
        section_break_idx = min(64, len(doc.paragraphs) - 1)
        
    return cover_end_idx, section_break_idx

def fill_document(template_path, content_data, output_path):
    print(f"Loading template: {template_path}...")
    doc = docx.Document(template_path)
    
    # 1. Dynamic template analysis
    cover_end_idx, section_break_idx = find_template_landmarks(doc)
    print(f"Landmarks found: Cover End Index = {cover_end_idx}, Section Break Index = {section_break_idx}")
    
    # 2. Fill Title Table (Table 0)
    if len(doc.tables) > 0 and "title" in content_data:
        title_table = doc.tables[0]
        # Set to Row 0, Cell 1
        set_cell_text(title_table.cell(0, 1), content_data["title"], font_name='宋体', font_size_pt=16, bold=True)
        print("Filled Title Table.")
        
    # 3. Fill Student Info Table (Table 1)
    if len(doc.tables) > 1 and "fields" in content_data:
        info_table = doc.tables[1]
        fields = content_data["fields"]
        for row in info_table.rows:
            label = row.cells[0].text.strip()
            # Clean label characters like "：" to match key
            cleaned_label = label.replace("：", "").strip()
            if cleaned_label in fields:
                val = fields[cleaned_label]
                set_cell_text(row.cells[2], val, font_name='宋体', font_size_pt=15, bold=False)
            elif row.cells[2].text.strip() == "小三宋体居中":
                # Clear placeholder if not provided
                set_cell_text(row.cells[2], "            ", font_name='宋体', font_size_pt=15, bold=False)
        print("Filled Student Info Table.")
        
    # 4. Clean paragraphs preserving the cover page and section break
    p_elements = list(doc.paragraphs)
    print(f"Original paragraph count: {len(p_elements)}")
    
    # Delete paragraphs after the section break (in reverse order)
    for idx in range(len(p_elements) - 1, section_break_idx, -1):
        p_elements[idx]._element.getparent().remove(p_elements[idx]._element)
        
    # Delete paragraphs between cover page and section break (in reverse order)
    for idx in range(section_break_idx - 1, cover_end_idx, -1):
        p_elements[idx]._element.getparent().remove(p_elements[idx]._element)
        
    # Refresh paragraph list
    p_sect_break = doc.paragraphs[cover_end_idx + 1]
    
    # 5. Write Table of Contents (before section break, in Section 0)
    if "toc" in content_data and content_data["toc"]:
        # Add Page Break before TOC to separate from cover
        p_pb = p_sect_break.insert_paragraph_before()
        p_pb.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        
        # TOC Title
        p_toc_title = p_sect_break.insert_paragraph_before()
        p_toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_toc_title.paragraph_format.space_before = Pt(0)
        p_toc_title.paragraph_format.space_after = Pt(20)
        p_toc_title.paragraph_format.line_spacing = Pt(20)
        p_toc_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run_toc_title = p_toc_title.add_run("目    录")
        set_run_font(run_toc_title, '黑体', 16, bold=True)
        
        for toc_item in content_data["toc"]:
            text = toc_item.get("text", "")
            page = toc_item.get("page", 1)
            level = toc_item.get("level", 0)
            add_toc_entry(p_sect_break, text, page, indent_level=level)
        print("Generated Table of Contents.")
        
    # 6. Append main body (after section break, in Section 1)
    if "body" in content_data:
        for idx, item in enumerate(content_data["body"]):
            el_type = item.get("type", "p")
            text = item.get("text", "")
            
            if el_type == "h1":
                add_heading_1(doc, text)
            elif el_type == "h2":
                add_heading_2(doc, text)
            elif el_type == "h3":
                add_heading_3(doc, text)
            elif el_type == "p":
                add_body_paragraph(doc, text)
            elif el_type == "ref_title":
                # References Title (黑体四号顶格)
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = Pt(20)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                run = p.add_run(text)
                set_run_font(run, '黑体', 14, bold=True)
            elif el_type == "ref_item":
                add_reference_item(doc, text)
            elif el_type == "app_title":
                # Appendix Title (黑体四号居中)
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = Pt(20)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                run = p.add_run(text)
                set_run_font(run, '黑体', 14, bold=True)
            elif el_type == "pb":
                doc.add_page_break()
                
        print("Generated Body text sections.")
        
    # Save document
    doc.save(output_path)
    print(f"Completed! Output saved to: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Automated Word Report Compiler based on docx template.")
    parser.add_argument("--template", required=True, help="Path to the docx template file")
    parser.add_argument("--content", required=True, help="Path to the JSON file containing report content")
    parser.add_argument("--output", required=True, help="Output path for the compiled docx document")
    args = parser.parse_args()
    
    if not os.path.exists(args.template):
        print(f"Error: Template file {args.template} not found.")
        exit(1)
    if not os.path.exists(args.content):
        print(f"Error: Content JSON file {args.content} not found.")
        exit(1)
        
    with open(args.content, 'r', encoding='utf-8') as f:
        content_data = json.load(f)
        
    fill_document(args.template, content_data, args.output)
