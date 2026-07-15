import os
import docx
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_run_font(run, font_name, font_size_pt, bold=False, italic=False):
    """
    Sets the font for a run. Sets the default/ASCII font to Times New Roman 
    and the East Asian font to the specified Chinese font (e.g., 宋体, 黑体, 楷体).
    """
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_cell_text(cell, text, font_name='宋体', font_size_pt=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """
    Sets the text of a table cell, clearing other paragraphs and preserving formatting and alignment.
    """
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(text)
    set_run_font(run, font_name, font_size_pt, bold=bold)
    
    # Delete extra paragraphs in cell
    for extra_p in list(cell.paragraphs[1:]):
        extra_p._element.getparent().remove(extra_p._element)

def add_toc_entry(p_ref, text, page_num, indent_level=0, line_spacing_pt=20):
    """
    Inserts a Table of Contents entry before the reference paragraph (p_ref).
    Automatically calculates dots to fill the row and aligns the page number.
    """
    p = p_ref.insert_paragraph_before()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if indent_level > 0:
        p.paragraph_format.left_indent = Pt(indent_level * 16)
        
    char_len = 0
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            char_len += 2
        else:
            char_len += 1
            
    dots_needed = 72 - char_len
    if dots_needed < 5:
        dots_needed = 5
    dots = "…" * (dots_needed // 2)
    
    is_bold = (indent_level == 0)
    
    run_text = p.add_run(text)
    set_run_font(run_text, '宋体', 12, bold=is_bold)
    
    run_dots = p.add_run(dots)
    set_run_font(run_dots, '宋体', 12, bold=False)
    
    run_num = p.add_run(str(page_num))
    set_run_font(run_num, '宋体', 12, bold=is_bold)

def add_heading_1(doc, text, line_spacing_pt=20):
    """
    Adds a Heading 1 (黑体四号 / 14pt Bold, left aligned, space before/after).
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '黑体', 14, bold=True)
    return p

def add_heading_2(doc, text, line_spacing_pt=20):
    """
    Adds a Heading 2 (黑体小四号 / 12pt Bold, left aligned, space before/after).
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '黑体', 12, bold=True)
    return p

def add_heading_3(doc, text, line_spacing_pt=20):
    """
    Adds a Heading 3 (楷体小四号 / 12pt Italic/Regular, left aligned).
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_run_font(run, '楷体', 12, bold=False)
    return p

def add_body_paragraph(doc, text, line_spacing_pt=20, indent_pt=24):
    """
    Adds a body text paragraph (宋体小四号 / 12pt, justified, 2-character indent).
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.first_line_indent = Pt(indent_pt)
    
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, bold=False)
    return p

def add_reference_item(doc, text, line_spacing_pt=20):
    """
    Adds a reference item (五号宋体 / 10.5pt, no indent).
    """
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    
    run = p.add_run(text)
    set_run_font(run, '宋体', 10.5, bold=False)
    return p

def clean_doc_keep_sections(doc, cover_end_idx, section_break_idx):
    """
    Cleans up paragraphs in the template docx to prepare for writing.
    Keeps paragraphs from 0 to cover_end_idx (e.g. 12) and keeps the section_break_idx paragraph.
    Deletes all others, ensuring section breaks remain intact.
    """
    p_elements = list(doc.paragraphs)
    
    # Delete paragraphs after the section break (in reverse order)
    for idx in range(len(p_elements) - 1, section_break_idx, -1):
        p_elements[idx]._element.getparent().remove(p_elements[idx]._element)
        
    # Delete paragraphs between cover page and section break (in reverse order)
    for idx in range(section_break_idx - 1, cover_end_idx, -1):
        p_elements[idx]._element.getparent().remove(p_elements[idx]._element)
        
    return doc.paragraphs[cover_end_idx + 1] # Return the paragraph holding the section break
